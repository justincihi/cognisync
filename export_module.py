"""
Export Module for Cognisync
Handles exporting session transcripts and analyses to various formats
"""

import os
import logging
from datetime import datetime
from flask import jsonify, send_file, request
from io import BytesIO

logger = logging.getLogger(__name__)

def add_export_routes(app, get_db, require_auth):
    """Add export routes to the Flask app"""
    
    @app.route('/api/sessions/<int:session_id>/export/<format_type>')
    @require_auth
    def export_session(session_id, format_type):
        """Export a session in the specified format (pdf, markdown, docx, txt)"""
        try:
            with get_db() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT * FROM therapy_sessions 
                    WHERE id = ? AND user_id = ?
                ''', (session_id, request.current_user['user_id']))
                session = cursor.fetchone()
                
                if not session:
                    return jsonify({'error': 'Session not found or access denied'}), 404
                
                # Prepare content
                content = format_session_content(session, format_type)
                
                # Generate filename
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                client_name = session['client_name'].replace(' ', '_')
                filename = f"{client_name}_{session['session_id']}.{format_type}"
                
                if format_type == 'txt' or format_type == 'markdown' or format_type == 'md':
                    # Return as text file
                    buffer = BytesIO(content.encode('utf-8'))
                    buffer.seek(0)
                    return send_file(
                        buffer,
                        as_attachment=True,
                        download_name=filename,
                        mimetype='text/plain' if format_type == 'txt' else 'text/markdown'
                    )
                elif format_type == 'pdf':
                    # Generate PDF
                    pdf_buffer = generate_pdf(content, session)
                    return send_file(
                        pdf_buffer,
                        as_attachment=True,
                        download_name=filename,
                        mimetype='application/pdf'
                    )
                elif format_type == 'docx':
                    # Generate DOCX
                    docx_buffer = generate_docx(content, session)
                    return send_file(
                        docx_buffer,
                        as_attachment=True,
                        download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                else:
                    return jsonify({'error': f'Unsupported format: {format_type}'}), 400
                    
        except Exception as e:
            logger.error(f"Export error: {e}")
            return jsonify({'error': 'Export failed'}), 500

    @app.route('/api/sessions/<int:session_id>/download')
    @require_auth
    def download_session_analysis(session_id):
        """Download session analysis as JSON"""
        try:
            with get_db() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT * FROM therapy_sessions 
                    WHERE id = ? AND user_id = ?
                ''', (session_id, request.current_user['user_id']))
                session = cursor.fetchone()
                
                if not session:
                    return jsonify({'error': 'Session not found or access denied'}), 404
                
                # Convert to dict
                session_data = dict(session)
                
                # Parse JSON fields
                import json
                if session_data.get('sentiment_analysis'):
                    try:
                        session_data['sentiment_analysis'] = json.loads(session_data['sentiment_analysis'])
                    except:
                        pass
                
                return jsonify({
                    'success': True,
                    'session': session_data
                })
                
        except Exception as e:
            logger.error(f"Download error: {e}")
            return jsonify({'error': 'Download failed'}), 500


def format_session_content(session, format_type):
    """Format session content for export"""
    
    if format_type in ['markdown', 'md']:
        return f"""# Clinical Session Analysis
        
## Session Information
- **Client:** {session['client_name']}
- **Session ID:** {session['session_id']}
- **Therapy Type:** {session['therapy_type']}
- **Format:** {session['summary_format']}
- **Date:** {session['created_at']}
- **Confidence Score:** {session['confidence_score']}%

## Transcript
{session['transcript'] or 'No transcript available'}

## Clinical Analysis
{session['analysis']}

## Sentiment Analysis
{session['sentiment_analysis'] or 'No sentiment analysis available'}

## Validation Notes
{session['validation_analysis'] or 'No validation notes'}

---
*Generated by Cognisync™ - AI-Powered Clinical Documentation*
*This document is for professional use only and should be reviewed by a licensed clinician.*
"""
    else:  # txt format
        return f"""CLINICAL SESSION ANALYSIS
========================

SESSION INFORMATION
------------------
Client: {session['client_name']}
Session ID: {session['session_id']}
Therapy Type: {session['therapy_type']}
Format: {session['summary_format']}
Date: {session['created_at']}
Confidence Score: {session['confidence_score']}%

TRANSCRIPT
----------
{session['transcript'] or 'No transcript available'}

CLINICAL ANALYSIS
----------------
{session['analysis']}

SENTIMENT ANALYSIS
-----------------
{session['sentiment_analysis'] or 'No sentiment analysis available'}

VALIDATION NOTES
---------------
{session['validation_analysis'] or 'No validation notes'}

---
Generated by Cognisync™ - AI-Powered Clinical Documentation
This document is for professional use only and should be reviewed by a licensed clinician.
"""


def generate_pdf(content, session):
    """Generate PDF from content"""
    try:
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Clinical Session Analysis", ln=True, align='C')
        pdf.ln(10)
        
        # Session info
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Session Information", ln=True)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 6, f"""Client: {session['client_name']}
Session ID: {session['session_id']}
Therapy Type: {session['therapy_type']}
Format: {session['summary_format']}
Date: {session['created_at']}
Confidence Score: {session['confidence_score']}%""")
        pdf.ln(5)
        
        # Analysis
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Clinical Analysis", ln=True)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 6, session['analysis'] or 'No analysis available')
        pdf.ln(5)
        
        # Footer
        pdf.set_font("Arial", "I", 8)
        pdf.multi_cell(0, 5, "Generated by Cognisync™ - AI-Powered Clinical Documentation\nThis document is for professional use only.")
        
        # Return as buffer
        buffer = BytesIO(pdf.output(dest='S').encode('latin-1'))
        buffer.seek(0)
        return buffer
        
    except ImportError:
        logger.warning("fpdf not installed, falling back to text export")
        # Fallback to text
        buffer = BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        # Fallback to text
        buffer = BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer


def generate_docx(content, session):
    """Generate DOCX from content"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Clinical Session Analysis', 0)
        title.alignment = 1  # Center
        
        # Session info
        doc.add_heading('Session Information', 1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Client', session['client_name']),
            ('Session ID', session['session_id']),
            ('Therapy Type', session['therapy_type']),
            ('Format', session['summary_format']),
            ('Date', session['created_at']),
            ('Confidence Score', f"{session['confidence_score']}%")
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
        
        # Analysis
        doc.add_heading('Clinical Analysis', 1)
        doc.add_paragraph(session['analysis'] or 'No analysis available')
        
        # Sentiment
        if session.get('sentiment_analysis'):
            doc.add_heading('Sentiment Analysis', 1)
            doc.add_paragraph(str(session['sentiment_analysis']))
        
        # Footer
        footer = doc.add_paragraph()
        footer_run = footer.add_run('\nGenerated by Cognisync™ - AI-Powered Clinical Documentation\nThis document is for professional use only.')
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        logger.warning("python-docx not installed, falling back to text export")
        # Fallback to text
        buffer = BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        # Fallback to text
        buffer = BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer

